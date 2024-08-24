import numpy as np
import seaborn as sns
from pandas import DataFrame
import matplotlib.pyplot as plt
from sklearn import tree
from jinja2 import Environment, FileSystemLoader,Template
import math
from jupyter_dash import JupyterDash
from sklearn.metrics import roc_curve
from dash import Dash, html, dcc, dash_table, callback
import plotly.express as px
import base64
import language_tool_python
import os
from docx import Document
from docx.shared import RGBColor
import pandas as pd
import shap
# Loading the folder that contains the txt templates

# file_loader = FileSystemLoader('./templates')
base_dir = '/content/drive/MyDrive/ColabNotebooks/DataQAHelperWithoutLLM/'
template_dir = os.path.join(base_dir, 'templates')
file_loader = FileSystemLoader(template_dir)

# Creating a Jinja Environment

env = Environment(loader=file_loader)

# Loading the Jinja templates from the folder
# For the regression
get_correlation = env.get_template('getcorrelation.txt')
model_comparison = env.get_template('modelcomparison.txt')
correlation_state = env.get_template('correlation.txt')
prediction_results = env.get_template('prediction.txt')
linearSummary = env.get_template('linearSummary.txt')
linearSummary2 = env.get_template('linearSummary2.txt')
linearA1_1 = env.get_template('linearA1-1.txt')
linearA1_2 = env.get_template('linearA1-2.txt')
linearA1_3 = env.get_template('linearA1-3.txt')
linearA2_1 = env.get_template('linearA2-1.txt')
linearA2_2 = env.get_template('linearA2-2.txt')
linearA2_3 = env.get_template('linearA2-3.txt')
linearSummary3 = env.get_template('linearSummary3.txt')
linearQuestion = env.get_template('linearQuestionset.txt')
regressionQuestion = env.get_template('regressionQuestionset.txt')
regressionSummary = env.get_template('regressionSummary.txt')
regressionvif = env.get_template('regression_vif.txt')
GBRA1_1 = env.get_template('GBRA1-1.txt')
GBRA1_2 = env.get_template('GBRA1-2.txt')
GBRA1_3 = env.get_template('GBRA1-3.txt')
GBRA2_1 = env.get_template('GBRA2-1.txt')
GBRA2_2 = env.get_template('GBRA2-2.txt')
GBRA2_3 = env.get_template('GBRA2-3.txt')
RFA1_1 = env.get_template('RFA1-1.txt')
RFA1_2 = env.get_template('RFA1-2.txt')
RFA1_3 = env.get_template('RFA1-3.txt')
RFA2_1 = env.get_template('RFA2-1.txt')
RFA2_2 = env.get_template('RFA2-2.txt')
RFA2_3 = env.get_template('RFA2-3.txt')
DTA1_1 = env.get_template('DTA1-1.txt')
DTA1_2 = env.get_template('DTA1-2.txt')
DTA1_3 = env.get_template('DTA1-3.txt')
DTA2_1 = env.get_template('DTA2-1.txt')
DTA2_2 = env.get_template('DTA2-2.txt')
DTA2_3 = env.get_template('DTA2-3.txt')
ENA1_1 = env.get_template('ENA1-1.txt')
ENA1_2 = env.get_template('ENA1-2.txt')
ENA1_3 = env.get_template('ENA1-3.txt')
LARA1_1= env.get_template('LARA1-1.txt')
LARA1_2= env.get_template('LARA1-2.txt')
LARA1_3= env.get_template('LARA1-3.txt')
ADAA1_1= env.get_template('ADAA1-1.txt')
ADAA1_2= env.get_template('ADAA1-2.txt')
ADAA1_3= env.get_template('ADAA1-3.txt')
KNRA1_1= env.get_template('KNRA1-1.txt')
KNRA1_2= env.get_template('KNRA1-2.txt')
KNRA1_3= env.get_template('KNRA1-3.txt')
RRA1_1= env.get_template('RRA1-1.txt')
RRA1_2= env.get_template('RRA1-2.txt')
RRA1_3= env.get_template('RRA1-3.txt')
LRA1_1= env.get_template('LRA1-1.txt')
LRA1_2= env.get_template('LRA1-2.txt')
LRA1_3= env.get_template('LRA1-3.txt')

R2evaluatemodel= env.get_template('R2evaluate_model.txt')
MAEevaluatemodel= env.get_template('MAEevaluate_model.txt')
CVR2evaluatemodel= env.get_template('CVR2evaluate_model.txt')

DecisionTree1 = env.get_template('decisiontree1.txt')
DecisionTree2 = env.get_template('decisiontree2.txt')
DecisionTree3 = env.get_template('decisiontree3.txt')
DecisionTreeQuestion = env.get_template('decisiontreequestion.txt')
gamStory = env.get_template('gamStory.txt')
GAMslinear_stats = env.get_template('GAMsLinearL1')
GAMslinear_R2 = env.get_template('GAMsLinearL2')
GAMslinear_P = env.get_template('GAMsLinearL3')
GAMslinear_sum = env.get_template('GAMsLinearL4')
GB1 = env.get_template('GB1')
GB2 = env.get_template('GB2')
GB3 = env.get_template('GB3')
piecewiseQuestion = env.get_template('piecewiseQuestion.txt')
piecewiseSummary = env.get_template('piecewiseSummary.txt')
piecewiseSummary2 = env.get_template('piecewiseSummary2.txt')
piecewiseSummary3 = env.get_template('piecewiseSummary3.txt')
piecewiseSummary4 = env.get_template('piecewiseSummary4.txt')
piecewiseSummary5 = env.get_template('piecewiseSummary5.txt')
piecewiseSummary6 = env.get_template('piecewiseSummary6.txt')
piecewiseSummary7 = env.get_template('piecewiseSummary7.txt')
pvaluesummary = env.get_template('pvaluesummary.txt')

#For Casestudy
piecewiseCP1=env.get_template('piecewiseCP1.txt')
piecewiseCP2=env.get_template('piecewiseCP2.txt')
piecewiseDRD1=env.get_template('piecewiseDRD1.txt')

# For the classifier
logisticSummary = env.get_template('logisticSummary.txt')
logisticSummary2 = env.get_template('logisticSummary2.txt')
logisticSummary3 = env.get_template('logisticSummary3.txt')
logisticQuestion = env.get_template('logisticQuestionset.txt')
LogA1_1 = env.get_template('LogA1-1.txt')
LogA1_2 = env.get_template('LogA1-2.txt')
LogA1_3 = env.get_template('LogA1-3.txt')
LogA2_1 = env.get_template('LogA2-1.txt')
LogA2_2 = env.get_template('LogA2-2.txt')
LogA2_3 = env.get_template('LogA2-3.txt')
LogA3_1 = env.get_template('LogA3-1.txt')
LogA3_2 = env.get_template('LogA3-2.txt')
LogA3_3 = env.get_template('LogA3-3.txt')
classifieraccuracy = env.get_template('classifierAccuracy.txt')
classifierauc = env.get_template('classifierAUC.txt')
classifiercv = env.get_template('classifierCvscore.txt')
classifierEvaluation=env.get_template('classifierEvaluation.txt')
classifierf1 = env.get_template('classifierF1.txt')
classifierprecision = env.get_template('classifierPrecision.txt')
classifierrecall = env.get_template('classifierRecall.txt')
classifiercoeff=   env.get_template('classifierCoeff.txt')
classifierimp = env.get_template('classifierImportant.txt')
ridgequestionset = env.get_template('ridgeQuestionset.txt')
ridgedecision = env.get_template('ridgeDecision.txt')
classifierquestionset = env.get_template('classifierQuestionset.txt')
# For the cluster
clusterQuestion = env.get_template('clusterQuestionset.txt')
clusterBestnum = env.get_template('clusterWSSS.txt')
clusterDaCa = env.get_template('clusterDaviesCalinski.txt')
clusterSil = env.get_template('clusterSilhouette.txt')
clusterGroup = env.get_template('clusterGroup.txt')

# For some basic function
basicdescription = env.get_template('basicdescription.txt')
simpletrend = env.get_template('simpletrend.txt')
modeloverfit = env.get_template('modeloverfit.txt')
correlation_story = env.get_template('correlation2.txt')
correlation_story_2 = env.get_template('correlation3.txt')
datasetinformation = env.get_template('dataset_information.txt')

# variables which each load a different segmented regression template
segmented_R2P = env.get_template('testPiecewisePwlfR2P')
segmented_R2 = env.get_template('testPiecewisePwlfR2')
segmented_P = env.get_template('testPiecewisePwlfP')
segmented_B = env.get_template('testPiecewisePwlfB')
segmented_GD1 = env.get_template('drugreport1')
segmented_GC1 = env.get_template('childreport1')

# For Aberdeen City CP
national_data_compare_story = env.get_template('national_data.txt')
register_story = env.get_template('register.txt')
risk_factor_story = env.get_template('risk_factor.txt')
reregister_story = env.get_template('reregister.txt')
remain_story = env.get_template('remain_story.txt')
enquiries_story = env.get_template('enquiries_story.txt')
timescales_story = env.get_template('CPtimescales.txt')
SCRA_story=env.get_template('SCRAdata.txt')

# For different dependent variables compared DRD
dc1 = env.get_template('dependentmagnificationcompare')
dc2 = env.get_template('samedependentmagnificationcompare')
dc3 = env.get_template('dependentquantitycompare')
dc4 = env.get_template('trendpercentagedescription')
dct = env.get_template('trenddescription')
tppc = env.get_template('twopointpeak_child')

# for different independent variables compared
idc1 = env.get_template('independentquantitycompare')
idtpc = env.get_template('independenttwopointcomparison')

# for batch processing
bp1 = env.get_template('batchprocessing1')
bp2 = env.get_template('batchprocessing2')

# for ChatGPT
databackground = env.get_template('databackground')
questionrequest = env.get_template('question_request.txt')

# for pycaret
automodelcompare1 = env.get_template('AMC1.txt')
automodelcompare2 = env.get_template('AMC2.txt')
pycaretimp = env.get_template('pycaret_imp.txt')
pycaretmodelfit = env.get_template('pycaret_modelfit.txt')
pycaretclassificationimp = env.get_template('pycaret_classificationimp.txt')
pycaretclassificationmodelfit = env.get_template('pycaret_classificationmodelfit.txt')
pycaretclassificationfit=env.get_template('pycaret_classificationfit.txt')
pycaretclassificoefimp1=env.get_template('coefimpsummary1.txt')
pycaretclassificoefimp2=env.get_template('coefimpsummary2.txt')
# for SKpipeline
pipeline_interpretation = env.get_template('pipeline_interpretation.txt')


def start_app():
    app_name = JupyterDash(__name__)
    listTabs = []
    return (app_name, listTabs)


def dash_tab_add(listTabs, label, child):
    listTabs.append(dcc.Tab(label=label, children=child))


def run_app(app_name, listTabs, portnum=8050):
    app_name.layout = html.Div([dcc.Tabs(listTabs)])
    app_name.run_server(mode='inline', debug=True, port=portnum)

def replace_placeholder_in_run(run, placeholder, replacement_text):
    if placeholder in run.text:
        run.text = run.text.replace(placeholder, replacement_text)

def replace_placeholder(paragraph, placeholder, replacement_text):
    for run in paragraph.runs:
        replace_placeholder_in_run(run, placeholder, replacement_text)

def reshape_data(data, colname):
    df = pd.DataFrame(data)

    if colname not in df.columns:
        raise ValueError(f"The specified column '{colname}' is not in the DataFrame")

    df_transposed = df.set_index(colname).T
    df_transposed.reset_index(inplace=True)
    df_transposed.rename(columns={'index': colname}, inplace=True)

    return df_transposed

def replace_placeholder_with_table(doc, placeholder, df):
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])

    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        run = hdr_cells[j].paragraphs[0].add_run(str(col))
        run.font.color.rgb = RGBColor(255, 0, 0)

    for i in range(df.shape[0]):
        row_cells = table.rows[i + 1].cells
        for j, value in enumerate(df.values[i]):
            run = row_cells[j].paragraphs[0].add_run(str(value))
            run.font.color.rgb = RGBColor(255, 0, 0)

    for para in doc.paragraphs:
        if placeholder in para.text:
            inline = para.runs
            for item in inline:
                if placeholder in item.text:
                    item.text = item.text.replace(placeholder, "")
                    move_table_after(table, para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        inline = para.runs
                        for item in inline:
                            if placeholder in item.text:
                                item.text = item.text.replace(placeholder, "")
                                move_table_after(table, para)


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._element
    p.addnext(tbl)


def replace_placeholder_with_color(paragraph, placeholder, replacement_text, color):
    if placeholder in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if placeholder in inline[i].text:
                text = inline[i].text.replace(placeholder, replacement_text)
                inline[i].text = text
                inline[i].font.color.rgb = color
                print(f"Replaced {placeholder} with {replacement_text} in paragraph.")

def fill_template(doc_path, output_path, replacements, color):
    doc = Document(doc_path)

    for para in doc.paragraphs:
        for placeholder, replacement_text in replacements.items():
            replace_placeholder_with_color(para, placeholder, replacement_text, color)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, replacement_text in replacements.items():
                        replace_placeholder_with_color(para, placeholder, replacement_text, color)

    doc.save(output_path)

def ThreeSubtab(child1,child2,child3):
    subtab1 = dcc.Tab(label='Simple Answer', children=[
        html.P(child1)
    ])

    subtab2 = dcc.Tab(label='Detailed Answer', children=[
        html.P(child2)
    ])

    subtab3 = dcc.Tab(label='More Detailed Answer', children=[
        html.P(child3)
    ])
    return (subtab1,subtab2,subtab3)

def ThreeSubQA(Q1,Q2,Q3,A1, A2, A3):
    subtab1 = dcc.Tab(label='Simple Question and Answer', children=[  html.B('Question: '),html.P(Q1),html.Br(),
        html.B('Answer: '),html.P(A1)
    ])

    subtab2 = dcc.Tab(label='Detailed Question and Answer', children=[html.B('Question: '),html.P(Q2),html.Br(),
                                                         html.B('Answer: '),
                                                         html.P(A2)
    ])

    subtab3 = dcc.Tab(label='More Detailed Question and Answer', children=[html.B('Question: '),html.P(Q3),html.Br(),
                                                              html.B('Answer: '),
                                                              html.P(A3)
    ])
    return (subtab1, subtab2, subtab3)

def TwoSubtab(child1, child2):
    subtab1 = dcc.Tab(label='Simple Answer', children=[
        html.P(child1)
    ])

    subtab2 = dcc.Tab(label='Detailed Answer', children=[
        html.P(child2)
    ])

    return (subtab1, subtab2)

def create_scatter_plots(data, Xcol, ycol):
    if not os.path.exists('pictures'):
        os.makedirs('pictures')
    for ind in Xcol:
        ax = sns.scatterplot(x=ind, y=ycol, data=data)
        plt.savefig(f'pictures/{ind}.png')
        plt.clf()

def TreeExplain(model, Xcol):
    n_nodes = model.tree_.node_count
    children_left = model.tree_.children_left
    children_right = model.tree_.children_right
    feature = model.tree_.feature
    threshold = model.tree_.threshold
    node_depth = np.zeros(shape=n_nodes, dtype=np.int64)
    is_leaves = np.zeros(shape=n_nodes, dtype=bool)
    stack = [(0, 0)]  # start with the root node id (0) and its depth (0)
    explain = ""
    while len(stack) > 0:
        # `pop` ensures each node is only visited once
        node_id, depth = stack.pop()
        node_depth[node_id] = depth
        # If the left and right child of a node is not the same we have a split
        # node
        is_split_node = children_left[node_id] != children_right[node_id]
        # If a split node, append left and right children and depth to `stack`
        # so we can loop through them
        if is_split_node:
            stack.append((children_left[node_id], depth + 1))
            stack.append((children_right[node_id], depth + 1))
        else:
            is_leaves[node_id] = True

    explain = explain + (
        "The binary tree structure has {n} nodes and has "
        "the following tree structure:\n ".format(n=n_nodes)
    )
    for i in range(n_nodes):
        if is_leaves[i]:
            explain = explain + (
                "{space}node={node} is a leaf node.\n".format(
                    space=node_depth[i] * "\t", node=i
                )
            )
        else:
            explain = explain + (
                "{space}node={node} is a split node: "
                "go to node {left} if {feature} <= {threshold:.2f} "
                "else to node {right}.\n".format(
                    space=node_depth[i] * "\t",
                    node=i,
                    left=children_left[i],
                    feature=Xcol[feature[i]],
                    threshold=threshold[i],
                    right=children_right[i],
                )
            )
    return (explain)


class Microplanning:
    def variablenamechange(self, dataset, Xcol, ycol, Xnewname, ynewname):
        if Xnewname != "":
            if np.size(Xnewname) != np.size(Xcol):
                raise Exception(
                    "The column name of the replacement X is inconsistent with the size of the column name of the original data X.")
            for i in range(np.size(Xnewname)):
                if (Xnewname[i] != ''):
                    dataset.rename(columns={Xcol[i]: Xnewname[i]}, inplace=True)
                else:
                    Xnewname[i] = Xcol[i]
        elif type(Xcol) == str and Xnewname == "":
            Xnewname = Xcol
        if (ynewname != ''):
            dataset.rename(columns={ycol: ynewname}, inplace=True)
        else:
            ynewname = ycol
        return (dataset, Xnewname, ynewname)

    def GrammarCorrection(self, text):
        tool = language_tool_python.LanguageTool('en-US')
        # get the matches
        matches = tool.check(text)
        my_mistakes = []
        my_corrections = []
        start_positions = []
        end_positions = []

        for rules in matches:
            if len(rules.replacements) > 0:
                start_positions.append(rules.offset)
                end_positions.append(rules.errorLength + rules.offset)
                my_mistakes.append(text[rules.offset:rules.errorLength + rules.offset])
                my_corrections.append(rules.replacements[0])

        my_new_text = list(text)

        for m in range(len(start_positions)):
            for i in range(len(text)):
                my_new_text[start_positions[m]] = my_corrections[m]
                if (i > start_positions[m] and i < end_positions[m]):
                    my_new_text[i] = ""
        my_new_text = "".join(my_new_text)
        return (my_new_text)



class DocumentplanningandDashboard:

    def LinearModelStats_view(self, linearData, r2, mape, mse, rmse, mae, vif, data, Xcol, ycol, questionset, expect,
                              portnum):

        if expect == "":
            expect = ["", "", ""]
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Store results for xcol
        for ind in linearData.index:
            ax = sns.regplot(x=ind, y=ycol, data=data)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()
        # Create Format index with file names
        _base64 = []
        for ind in linearData.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))

        linear_app, listTabs = start_app()

        # Add to dashbord Linear Model Statistics
        fig = px.bar(linearData)

        intro = linearSummary2.render(r2=r2, indeNum=np.size(Xcol), modelName="Linear Model", Xcol=Xcol,
                                      ycol=ycol, qs=questionset, t=expect[0], expect=expect[1])

        aim = Xcol
        aim.insert(0, ycol)

        question1 = linearQuestion.render(xcol=Xcol, ycol=ycol, qs=questionset, section=1, indeNum=np.size(Xcol),
                                         trend=expect[0])
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        text1 = intro+' '+linearA1_1.render(r2=r2)
        text2 = intro+' '+linearA1_2.render(r2=r2, mape=mape)
        text3 = intro+' '+linearA1_3.render(r2=r2, mape=mape, mse=mse, rmse=rmse, mae=mae)

        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # add subtab to the tab
        children = [ dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in data[aim].columns],
                                         style_table={'height': '400px', 'overflowY': 'auto'}),
                    dcc.Tabs([subtab1, subtab2, subtab3])
                    ]

        dash_tab_add(listTabs, 'LinearModelStats', children)
        aim.remove(ycol)

        pf, nf, nss, ss, imp, i = "", "", "", "", "", 0
        # Add to dashbord Xcol plots and data story

        for ind in linearData.index:
            question1 = linearQuestion.render(xcol=ind, ycol=ycol, qs=[0,0,1], section=2, indeNum=1, trend=expect[0])
            question2 = linearQuestion.render(xcol=ind, ycol=ycol, qs=questionset, section=2, indeNum=1, trend=expect[0])
            question3=question2+' '+linearQuestion.render(section=7)
            text1 = linearA2_1.render(xcol=ind, ycol=ycol, coeff=linearData['coeff'][ind])
            text2 = linearA2_2.render(xcol=ind, ycol=ycol, coeff=linearData['coeff'][ind],
                                      p=linearData['pvalue'][ind])
            text3 = linearA2_3.render(xcol=ind, ycol=ycol, coeff=linearData['coeff'][ind],
                                      p=linearData['pvalue'][ind],
                                      vif_value=vif[vif['feature'] == ind]['VIF'].values[0])

            subtab1, subtab2, subtab3 = ThreeSubQA(question1, question2, question3,text1,text2,text3)

            # newstory = MicroLexicalization(story)
            if abs(linearData['coeff'][ind]) == max(abs(linearData['coeff'])):
                imp = ind
            if linearData['coeff'][ind] > 0:
                pf = pf + "the " + ind + ", "
            elif linearData['coeff'][ind] < 0:
                nf = nf + "the " + ind + ", "
            if linearData['pvalue'][ind] > 0.05:
                nss = nss + "the " + ind + ", "
            else:
                ss = ss + "the " + ind + ", "

            if questionset[1] == 1 or questionset[2] == 1:
                children = [
                    html.Img(src='data:image/png;base64,{}'.format(_base64[i])),
                    dcc.Tabs([subtab1, subtab2, subtab3])
                ]
                dash_tab_add(listTabs, ind, children)

            i = i + 1
        question = linearQuestion.render(xcol="", ycol=ycol, qs=questionset, section=3, indeNum=1, trend=expect[0])
        summary = linearSummary3.render(imp=imp, ycol=ycol, nss=nss, ss=ss, pf=pf, nf=nf, t=expect[0], r2=r2,
                                        qs=questionset)

        children = [dcc.Graph(figure=fig), html.B('Question: '),html.P(question), html.Br(), html.B('Answer: '),html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(linear_app, listTabs, portnum)

    def LogisticModelStats_view(self, model, data, Xcol, ycol, devDdf, accuracy, auc, pos_class_mean, questionset,
                                portnum):

        # Create data frames for coefficients, p-values and importance scores.
        columns1 = {'coeff': model.params.values, 'pvalue': model.pvalues.round(4).values}
        logisticData1 = DataFrame(data=columns1, index=Xcol)
        columns2 = {'importance score': abs(model.params.values)}
        logisticData2 = DataFrame(data=columns2, index=Xcol)
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Store results for xcol
        for ind in logisticData1.index:
            ax = sns.regplot(x=ind, y=ycol, data=data, logistic=True)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()

        # Create Format index with file names
        _base64 = []

        for ind in logisticData1.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))
        logistic_app, listTabs = start_app()
        i = 0

        # Add to dashbord Model Statistics
        question = logisticQuestion.render(indeNum=np.size(Xcol), xcol=Xcol, ycol=ycol, qs=questionset, section=1)
        intro = logisticSummary3.render(indeNum=np.size(Xcol), modelName="logistic regression model", Xcol=Xcol,
                                        ycol=ycol)
        aim = Xcol
        aim.insert(0, ycol)

        text1 = LogA1_1.render(modelName="logistic regression model", accuracy=accuracy)
        text2 = LogA1_2.render(modelName="logistic regression model", accuracy=accuracy, auc=auc)
        text3 = LogA1_3.render(modelName="logistic regression model", accuracy=accuracy, auc=auc, ddd=devDdf)

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        # add subtab to the tab
        children = [html.P(question), html.Br(), html.P(intro),
                    dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in data[aim].columns],
                                         style_table={'height': '400px', 'overflowY': 'auto'})]

        dash_tab_add(listTabs, 'LogisticModelStats', children)
        aim.remove(ycol)

        pos_eff, neg_eff, nss, ss, imp = "", "", "", "", ""
        # Add to dashbord Xcol plots and data story

        for ind in logisticData1.index:
            question = logisticQuestion.render(indeNum=1, xcol=ind, ycol=ycol, qs=questionset, section=2)
            # independent_variable_story
            # independent_variable_story = logisticSummary.render(xcol=ind, ycol=ycol,
            #                                                     odd=abs(
            #                                                         100 * (math.exp(logisticData1['coeff'][ind]) - 1)),
            #                                                     coeff=logisticData1['coeff'][ind],
            #                                                     p=logisticData1['pvalue'][ind],
            #                                                     qs=questionset)

            text1 = LogA2_1.render(xcol=ind, coeff=logisticData1['coeff'][ind], pos_class_mean=pos_class_mean)
            text2 = LogA2_2.render(xcol=ind, odd=abs(
                100 * (math.exp(logisticData1['coeff'][ind]) - 1)),
                                   coeff=logisticData1['coeff'][ind], pos_class_mean=pos_class_mean)
            text3 = LogA2_3.render(xcol=ind,
                                   odd=abs(
                                       100 * (math.exp(logisticData1['coeff'][ind]) - 1)),
                                   coeff=logisticData1['coeff'][ind],
                                   p=logisticData1['pvalue'][ind], pos_class_mean=pos_class_mean
                                   )

            subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)
            # independent_variable_story = model.MicroLexicalization(independent_variable_story)
            if logisticData1['coeff'][ind] == max(logisticData1['coeff']):
                imp = ind
            if logisticData1['coeff'][ind] > 0:
                pos_eff = pos_eff + ind + ", "
            else:
                neg_eff = neg_eff + ind + ", "
            if logisticData1['pvalue'][ind] > 0.05:
                nss = nss + ind + ", "
            else:
                ss = ss + ind + ", "
            if questionset[1] == 1 or questionset[2] == 1:
                children = [
                    html.Img(src='data:image/png;base64,{}'.format(_base64[i])), html.P(question), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])
                ]
                dash_tab_add(listTabs, ind, children)
            i = i + 1
        fig = px.bar(logisticData1)

        filtered_df = logisticData1[logisticData1['pvalue'] < 0.05]
        impWithP = filtered_df['coeff'].abs().idxmax()

        plt.savefig('pictures/{}.png'.format(imp))
        plt.clf()
        question = logisticQuestion.render(indeNum=1, xcol=ind, ycol=ycol, qs=questionset, section=3)
        # summary = model.MicroLexicalization(summary)
        # summary = LogA3_1.render(modelName="logistic regression model",accuracy=accuracy,pos=pos_eff, neg=neg_eff, nss=nss, ss=ss, imp=imp)

        text1 = LogA3_1.render(modelName="logistic regression model", accuracy=accuracy, pos=pos_eff, neg=neg_eff,
                               nss=nss, ss=ss, imp=imp)
        text2 = LogA3_2.render(modelName="logistic regression model", accuracy=accuracy, auc=auc, pos=pos_eff,
                               neg=neg_eff, nss=nss, ss=ss, imp=imp)
        text3 = LogA3_3.render(modelName="logistic regression model", accuracy=accuracy, pos=pos_eff, auc=auc,
                               ddd=devDdf, neg=neg_eff, nss=nss, ss=ss, imp=imp, impwithp=impWithP)

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)
        children = [
            dcc.Graph(figure=fig), html.P(question), html.Br(),
            dcc.Tabs([subtab1, subtab2, subtab3])
        ]
        # children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(logistic_app, listTabs, portnum)

    def GradientBoostingModelStats_view(self, data, Xcol, ycol, GBmodel, mse, mae, r2, imp, lessimp, questionset,
                                        gbr_params, train_r2, test_r2, train_mae, test_mae, cv_r2_scores, cv_r2_mean,
                                        mape, imp_pos_ave, imp_pos_value_ave, imp_neg_ave, imp_neg_value_ave,imp_var,explainer,shap_values,X_test,portnum):

        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Store importance figure
        plt.bar(Xcol, GBmodel.feature_importances_)

        plt.title("Importance Score")
        plt.savefig('pictures/{}.png'.format("GB1"))
        plt.clf()

        _base64 = []
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("GB1"), 'rb').read()).decode('ascii'))
        # Training & Test Deviance Figure
        test_score = np.zeros((gbr_params['n_estimators'],), dtype=np.float64)
        fig = plt.figure(figsize=(8, 8))
        plt.title('Deviance')
        plt.plot(np.arange(gbr_params['n_estimators']) + 1, GBmodel.train_score_, 'b-',
                 label='Training Set Deviance')
        plt.plot(np.arange(gbr_params['n_estimators']) + 1, test_score, 'r-',
                 label='Test Set Deviance')
        plt.legend(loc='upper right')
        plt.xlabel('Boosting Iterations')
        plt.ylabel('Deviance')
        fig.tight_layout()
        plt.savefig('pictures/{}.png'.format("GB2"))
        plt.clf()
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("GB2"), 'rb').read()).decode('ascii'))

        plt.plot(cv_r2_scores, label='Cross-Validation R²')
        plt.axhline(y=cv_r2_mean, color='r', linestyle='--', label='Mean Cross-Validation R²')
        plt.xlabel('Fold')
        plt.ylabel('R²')
        plt.title('Cross-Validation R² Scores')
        plt.legend()
        plt.savefig('pictures/{}.png'.format("GB3"))
        plt.clf()
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("GB3"), 'rb').read()).decode('ascii'))

        GB_app, listTabs = start_app()
        # Add to dashbord Model Statistics

        text1 = GBRA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = GBRA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = GBRA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)
        # micro planning
        # intro = model.MicroLexicalization(intro)

        question1 = DecisionTreeQuestion.render(q=1, m="gb")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'Gradient Boosting Stats', children)
        aim.remove(ycol)

        text1 = GBRA2_1.render(imp=imp)
        text2 = GBRA2_2.render(imp=imp, lessimp=lessimp)
        feature_importances = GBmodel.feature_importances_
        importance_df = pd.DataFrame({'Feature': Xcol, 'Importance': feature_importances})
        importance_df = importance_df.sort_values(by='Importance', ascending=False).reset_index(drop=True)
        sorted_Xcol = importance_df['Feature'].tolist()
        text3 = GBRA2_3.render(sorted_Xcol=sorted_Xcol)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        question3 = DecisionTreeQuestion.render(q=3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question3), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Feature Importances', children)

        question4 = DecisionTreeQuestion.render(q=4)
        text1 = CVR2evaluatemodel.render(cv_r2_mean=cv_r2_mean, train_r2=train_r2, diff=abs(train_r2 - cv_r2_mean))
        text2 = text1 + '\n' + R2evaluatemodel.render(train_r2=train_r2, test_r2=test_r2, diff=abs(train_r2 - test_r2))
        text3 = text2 + '\n' + MAEevaluatemodel.render(train_mae=train_mae, test_mae=test_mae,
                                                       diff=abs(train_mae - test_mae))

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question4), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Model Fitting', children)


        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_values, X_test, feature_names=Xcol, show=False)
        plt.savefig('pictures/GB3.png', bbox_inches='tight')
        plt.close()

        _base64.append(base64.b64encode(open('pictures/{}.png'.format("GB3"), 'rb').read()).decode('ascii'))

        question=regressionQuestion.render(section=4)
        shapStory = pycaretimp.render(imp=imp_var, target=ycol, imp_pos_ave=imp_pos_ave,
                                   imp_pos_value_ave=imp_pos_value_ave,
                                   imp_neg_ave=imp_neg_ave, imp_neg_value_ave=imp_neg_value_ave)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[3])), html.P(question), html.Br(),html.P(shapStory)]
        dash_tab_add(listTabs, 'SHAP Interpretation', children)

        run_app(GB_app, listTabs, portnum)

    def RandomForestRegressionModelStats_view(self, data, Xcol, ycol, rf_small, DTData, r2, mse, mae, mape, train_r2,
                                              test_r2, train_mae, test_mae, cv_r2_scores, cv_r2_mean,imp_pos_ave, imp_pos_value_ave, imp_neg_ave, imp_neg_value_ave,imp_var, explainer,shap_values,X_test,portnum=8050):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')

        plt.bar(Xcol, rf_small.feature_importances_)

        plt.title("Importance Score")
        plt.savefig('pictures/{}.png'.format("RF1"))
        plt.clf()

        _base64 = []
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("RF1"), 'rb').read()).decode('ascii'))

        plt.plot(cv_r2_scores, label='Cross-Validation R²')
        plt.axhline(y=cv_r2_mean, color='r', linestyle='--', label='Mean Cross-Validation R²')
        plt.xlabel('Fold')
        plt.ylabel('R²')
        plt.title('Cross-Validation R² Scores')
        plt.legend()
        plt.savefig('pictures/{}.png'.format("RF2"))
        plt.clf()
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("RF2"), 'rb').read()).decode('ascii'))

        # Importance score Figure
        imp, lessimp = "", ""
        for ind in DTData.index:
            if DTData['important'][ind] == max(DTData['important']):
                imp = ind
            elif DTData['important'][ind] == min(DTData['important']):
                lessimp = ind
        # print(DTData['important'])
        RF_app, listTabs = start_app()
        # Add to dashbord Model Statistics



        text1 = RFA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = RFA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = RFA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                              mae=mae)
        # micro planning
        # intro = model.MicroLexicalization(intro)
        question1 = DecisionTreeQuestion.render(q=1, m="rf")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'RandomForestModelStats', children)

        aim.remove(ycol)

        text1 = RFA2_1.render(imp=imp)
        text2 = RFA2_2.render(imp=imp, lessimp=lessimp)
        feature_importances = rf_small.feature_importances_
        importance_df = pd.DataFrame({'Feature': Xcol, 'Importance': feature_importances})
        importance_df = importance_df.sort_values(by='Importance', ascending=False).reset_index(drop=True)
        sorted_Xcol = importance_df['Feature'].tolist()
        text3 = RFA2_3.render(sorted_Xcol=sorted_Xcol)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        question3 = DecisionTreeQuestion.render(q=3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question3), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Feature Importances', children)

        question4 = DecisionTreeQuestion.render(q=4)
        text1 = CVR2evaluatemodel.render(cv_r2_mean=cv_r2_mean, train_r2=train_r2, diff=abs(train_r2 - cv_r2_mean))
        text2 = text1 + '\n' + R2evaluatemodel.render(train_r2=train_r2, test_r2=test_r2, diff=abs(train_r2 - test_r2))
        text3 = text2 + '\n' + MAEevaluatemodel.render(train_mae=train_mae, test_mae=test_mae,
                                                       diff=abs(train_mae - test_mae))

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question4), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Model Fitting', children)

        shap.initjs()
        shap.force_plot(explainer.expected_value, shap_values[0, :], X_test[0, :], feature_names=Xcol)

        # Summary plot for the whole dataset

        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_values, X_test, feature_names=Xcol, show=False)
        plt.savefig('pictures/RF3.png', bbox_inches='tight')
        plt.close()

        _base64.append(base64.b64encode(open('pictures/{}.png'.format("RF3"), 'rb').read()).decode('ascii'))

        question=regressionQuestion.render(section=4)
        shapStory = pycaretimp.render(imp=imp_var, target=ycol, imp_pos_ave=imp_pos_ave,
                                   imp_pos_value_ave=imp_pos_value_ave,
                                   imp_neg_ave=imp_neg_ave, imp_neg_value_ave=imp_neg_value_ave)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),html.P(shapStory)]
        dash_tab_add(listTabs, 'SHAP Interpretation', children)

        run_app(RF_app, listTabs, portnum)

    def DecisionTreeRegressionModelStats_view(self, data, Xcol, ycol, DTData, DTmodel, r2, mse, mae, mape, train_r2,
                                              test_r2, train_mae, test_mae, cv_r2_scores, cv_r2_mean, imp_pos_ave, imp_pos_value_ave, imp_neg_ave, imp_neg_value_ave,imp_var,explainer,shap_values,X_test,
                                              portnum=8050):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Importance score Figure
        imp = ""
        plt.bar(Xcol, DTmodel.feature_importances_)

        plt.title("Importance Score")
        plt.savefig('pictures/{}.png'.format("DT1"))
        plt.clf()

        _base64 = []
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("DT1"), 'rb').read()).decode('ascii'))

        plt.plot(cv_r2_scores, label='Cross-Validation R²')
        plt.axhline(y=cv_r2_mean, color='r', linestyle='--', label='Mean Cross-Validation R²')
        plt.xlabel('Fold')
        plt.ylabel('R²')
        plt.title('Cross-Validation R² Scores')
        plt.legend()
        plt.savefig('pictures/{}.png'.format("DT2"))
        plt.clf()
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("DT2"), 'rb').read()).decode('ascii'))

        imp, lessimp = "", ""
        for ind in DTData.index:
            if DTData['important'][ind] == max(DTData['important']):
                imp = ind
            elif DTData['important'][ind] == min(DTData['important']):
                lessimp = ind
        DT_app, listTabs = start_app()

        # Add to dashbord Model Statistics
        text1 = DTA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = DTA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = DTA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                              mae=mae)
        # micro planning
        # intro = model.MicroLexicalization(intro)

        question1 = DecisionTreeQuestion.render(q=1, m="dt")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'DecisionTreeModelStats', children)
        aim.remove(ycol)

        text1 = DTA2_1.render(imp=imp)
        text2 = DTA2_2.render(imp=imp, lessimp=lessimp)
        feature_importances = DTmodel.feature_importances_
        importance_df = pd.DataFrame({'Feature': Xcol, 'Importance': feature_importances})
        importance_df = importance_df.sort_values(by='Importance', ascending=False).reset_index(drop=True)
        sorted_Xcol = importance_df['Feature'].tolist()
        text3 = DTA2_3.render(sorted_Xcol=sorted_Xcol)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        question3 = DecisionTreeQuestion.render(q=3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question3), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Feature Importances', children)

        question4 = DecisionTreeQuestion.render(q=4)
        text1 = CVR2evaluatemodel.render(cv_r2_mean=cv_r2_mean, train_r2=train_r2, diff=abs(train_r2 - cv_r2_mean))
        text2 = text1 + '\n' + R2evaluatemodel.render(train_r2=train_r2, test_r2=test_r2, diff=abs(train_r2 - test_r2))
        text3 = text2 + '\n' + MAEevaluatemodel.render(train_mae=train_mae, test_mae=test_mae,
                                                       diff=abs(train_mae - test_mae))

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question4), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Model Fitting', children)

        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_values, X_test, feature_names=Xcol, show=False)
        plt.savefig('pictures/DT3.png', bbox_inches='tight')
        plt.close()

        _base64.append(base64.b64encode(open('pictures/{}.png'.format("DT3"), 'rb').read()).decode('ascii'))

        question=regressionQuestion.render(section=4)
        shapStory = pycaretimp.render(imp=imp_var, target=ycol, imp_pos_ave=imp_pos_ave,
                                   imp_pos_value_ave=imp_pos_value_ave,
                                   imp_neg_ave=imp_neg_ave, imp_neg_value_ave=imp_neg_value_ave)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),html.P(shapStory)]
        dash_tab_add(listTabs, 'SHAP Interpretation', children)


        # Figure of the tree
        fig2, ax = plt.subplots(figsize=(20, 10), dpi=200)
        tree.plot_tree(DTmodel,
                       feature_names=Xcol,
                       class_names=ycol,
                       filled=True,
                       node_ids=True);
        fig2.savefig('pictures/{}.png'.format("DT"))
        plt.clf()
        encoded_image = base64.b64encode(open("pictures/DT.png", 'rb').read()).decode('ascii')

        # # Explain of the tree
        explain = TreeExplain(DTmodel, Xcol)
        # Text need to fix here
        tree_explain_story = explain
        question2 = DecisionTreeQuestion.render(q=2)
        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image)),
                    html.P(question2), html.Br(), html.Pre(tree_explain_story)]
        dash_tab_add(listTabs, 'Tree Explanation', children)


        run_app(DT_app, listTabs, portnum)

    def piecewise_linear_view(self, data, Xcol, ycol, model, slopes, segment_points, segment_values, max_slope_info,
                              breaks, segment_r2_values, mse, mae, bic, aic, portnum):

        x_hat = np.linspace(data[Xcol].min(), data[Xcol].max(), 100)
        y_hat = model.predict(x_hat)
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        plt.figure(figsize=(10, 6))
        sns.scatterplot(x=Xcol, y=ycol, data=data)
        plt.plot(x_hat, y_hat, label='Piecewise Linear Fit', color='red')
        plt.xlabel(Xcol)
        plt.ylabel(ycol)
        plt.legend()
        plt.title('Piecewise Linear Fit')
        plt.savefig('pictures/piecewise_linear_fit.png')
        encoded_image = base64.b64encode(open("pictures/piecewise_linear_fit.png", 'rb').read()).decode('ascii')
        r2 = model.r_squared()

        plt.clf()
        sns.scatterplot(x=Xcol, y=ycol, data=data)
        for i in range(len(breaks) - 1):
            start, end = breaks[i], breaks[i + 1]
            x_segment = np.linspace(start, end, 100)
            y_segment = model.predict(x_segment)
            plt.plot(x_segment, y_segment, label=f'Segment {i + 1}', linestyle='--')

        plt.xlabel(Xcol)
        plt.ylabel(ycol)
        plt.legend()
        plt.title('Piecewise Linear Fit with Individual Segments')
        plt.savefig('pictures/piecewise_linear_fit_with_segments.png')
        plt.clf()
        encoded_image2 = base64.b64encode(open("pictures/piecewise_linear_fit_with_segments.png", 'rb').read()).decode(
            'ascii')

        piecewise_app, listTabs = start_app()

        def merge_segments(breaks, slopes):
            merged_breaks = [breaks[0]]
            merged_slopes = []
            current_slope = slopes[0]
            current_start = breaks[0]
            for i in range(1, len(slopes)):
                if (slopes[i] > 0 and current_slope > 0) or (slopes[i] < 0 and current_slope < 0) or (
                        slopes[i] == 0 and current_slope == 0):
                    continue
                else:
                    merged_breaks.append(breaks[i])
                    merged_slopes.append(current_slope)
                    current_slope = slopes[i]
                    current_start = breaks[i]
            merged_breaks.append(breaks[-1])
            merged_slopes.append(current_slope)
            output_slopes = [1 if s > 0 else -1 if s < 0 else 0 for s in merged_slopes]
            return merged_breaks, output_slopes

        merged_breaks, output_slopes = merge_segments(breaks, slopes)

        question = piecewiseQuestion.render(xcol=Xcol, ycol=ycol, section=4)
        summary = piecewiseSummary4.render(xcol=Xcol, ycol=ycol, merged_breaks=merged_breaks,
                                           output_slopes=output_slopes,
                                           startPoint=max_slope_info['start_end_points'][0],
                                           endPoint=max_slope_info['start_end_points'][1],
                                           slope=max_slope_info['slope'])

        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image)), html.P(question), html.Br(),
                    html.P(summary)]
        dash_tab_add(listTabs, 'Overview', children)

        question = piecewiseQuestion.render(xcol=Xcol, ycol=ycol, section=2)
        summary = ""
        for i in range(len(slopes)):
            summary = summary + piecewiseSummary2.render(startPoint=segment_points[i][0], endPoint=segment_points[i][1],
                                                         slope=slopes[i]) + "\n"
        text2 = piecewiseSummary5.render(breaks=breaks) + summary

        subtab1, subtab2 = TwoSubtab(summary, text2)

        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image2)), html.P(question), html.Br(),
                    dcc.Tabs([subtab1, subtab2])]
        dash_tab_add(listTabs, 'Slopes', children)

        question = piecewiseQuestion.render(xcol=Xcol, ycol=ycol, section=1)
        text1 = piecewiseSummary.render(r2=r2, modelName="Piecewise Linear Model")
        rmse = mse ** (1 / 2.0)
        text2 = text1 + piecewiseSummary7.render(AIC=aic, BIC=bic)
        text3 = text2 + piecewiseSummary6.render(mse=mse, rmse=rmse, mae=mae)

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image)), html.P(question), html.Br(),
                    dcc.Tabs([subtab1, subtab2, subtab3])]
        dash_tab_add(listTabs, 'Model Information', children)

        question = piecewiseQuestion.render(xcol="", ycol=ycol, section=3)
        summary = piecewiseSummary3.render(xcol=Xcol, ycol=ycol, startPoint=max_slope_info['start_end_points'][0],
                                           endPoint=max_slope_info['start_end_points'][1],
                                           slope=max_slope_info['slope'], yValue1=max_slope_info['start_end_values'][0],
                                           yValue2=max_slope_info['start_end_values'][1])

        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image2)), html.P(question), html.Br(),
                    html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(piecewise_app, listTabs, portnum)

    def RidgeClassifier_view(self, data, Xcol, ycol, rclf, pca, y_test, y_prob, roc_auc, X_pca, accuracy, precision,
                             recall, f1, importances,
                             class1,
                             class2, confusionmatrix, cv_scores):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        _base64 = []
        ridge_app, listTabs = start_app()
        # Plot ROC curve
        fpr, tpr, thresholds = roc_curve(y_test, y_prob)
        plt.figure(figsize=(8, 6))
        plt.plot(fpr, tpr, label='ROC Curve (area = {:.2f})'.format(roc_auc))
        plt.plot([0, 1], [0, 1], 'k--', label='Random')
        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        plt.title('Receiver Operating Characteristic')
        plt.legend(loc='lower right')
        plt.savefig('pictures/{}.png'.format("ROC"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("ROC"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot decision boundaries
        sns.scatterplot(x=X_pca[:, 0], y=X_pca[:, 1], hue=y_test, palette='husl')
        ax = plt.gca()
        xlim = ax.get_xlim()
        ylim = ax.get_ylim()

        xx, yy = np.meshgrid(np.linspace(xlim[0], xlim[1], 500),
                             np.linspace(ylim[0], ylim[1], 500))
        Z = rclf.predict(pca.inverse_transform(np.c_[xx.ravel(), yy.ravel()])).reshape(xx.shape)
        ax.contourf(xx, yy, Z, alpha=0.8, cmap='coolwarm')
        plt.xlabel('Principal Component 1')
        plt.ylabel('Principal Component 2')
        # plt.title('Decision Boundaries (Malignant vs. Benign)')
        plt.title('Decision Boundaries')
        plt.savefig('pictures/{}.png'.format("DecisionBoundaries"))
        _base64.append(
            base64.b64encode(open('pictures/{}.png'.format("DecisionBoundaries"), 'rb').read()).decode('ascii'))

        # Create a dataframe to store feature importances along with their corresponding feature names
        X = data[Xcol]
        importances_df = DataFrame({'Feature': X.columns, 'Importance': importances})
        # Sort the dataframe by importance in descending order
        importances_df = importances_df.sort_values(by='Importance', ascending=False)
        # Plot feature importances using a bar plot
        plt.figure(figsize=(10, 6))
        plt.bar(importances_df['Feature'], importances_df['Importance'])
        plt.xlabel('Feature')
        plt.ylabel('Importance')
        plt.title('Feature Importances')
        plt.xticks(rotation=90)
        plt.savefig('pictures/{}.png'.format("FeatureImportances"))

        # Plot confusion matrix
        sns.heatmap(confusionmatrix, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.title('Confusion Matrix')
        plt.savefig('pictures/{}.png'.format("confusionmatrix"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("confusionmatrix"), 'rb').read()).decode('ascii'))
        plt.clf()

        _base64.append(
            base64.b64encode(open('pictures/{}.png'.format("FeatureImportances"), 'rb').read()).decode('ascii'))
        for i, importance in enumerate(importances):
            if abs(importance) == max(abs(importances)):
                # print("Feature {}: {}".format(X.columns[i], importance))
                imp = X.columns[i]

        # Extract coefficients from the trained model
        coefs = rclf.coef_[0]
        intercept = rclf.intercept_[0]

        # Construct the linear equation for the decision boundaries
        equation = 'Decision Boundary Equation: '
        for i in range(len(Xcol)):
            equation += '({:.4f} * {}) + '.format(coefs[i], Xcol[i])
        equation += '{:.4f}'.format(intercept)

        text1 = classifieraccuracy.render(accuracy=round(accuracy, 3), classifiername="ridge classifier")
        text2 = text1 + classifierprecision.render(precision=precision)
        text3 = text2 + classifierrecall.render(recall=recall) + classifierf1.render(f1=f1)

        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)
        # dcc.Tabs([subtab1, subtab2, subtab3])

        question = ridgequestionset.render(section=1)
        aim = Xcol
        aim.insert(0, ycol)
        children = [html.P(question), html.Br(), dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'})]

        dash_tab_add(listTabs, 'RidgeClassifierStats', children)
        aim.remove(ycol)

        question = ridgequestionset.render(section=2)
        DecisionBoundaryStory = ridgedecision.render(equation=equation, class1=class1, class2=class2)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question), html.Br(),
                    html.P(DecisionBoundaryStory)]
        dash_tab_add(listTabs, "Decision Boundary Equation", children)

        question = ridgequestionset.render(section=3)
        aucStory = classifierauc.render(AUC=roc_auc)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question), html.Br(),
                    html.P(aucStory)]
        dash_tab_add(listTabs, "Area under the Receiver Operating Characteristic curve", children)

        question = classifierquestionset.render(section=3)
        crossvalidStory = classifiercv.render(cv=round(cv_scores.mean(), 3), cm=confusionmatrix)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),
                    html.P(crossvalidStory)]
        dash_tab_add(listTabs, "Confusion Matrix and Cross-validation", children)

        question = ridgequestionset.render(section=4)
        ImpStory = classifierimp.render(imp=imp)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[3])), html.P(question), html.Br(),
                    html.P(ImpStory)]
        dash_tab_add(listTabs, "Feature Importances", children)

        run_app(ridge_app, listTabs)

    def KNeighborsClassifier_view(self, data, Xcol, ycol, accuracy, precision, feature_importances, recall, f1,
                                  confusionmatrix, cv_scores):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        _base64 = []
        KNei_app, listTabs = start_app()
        # Print feature importances with column names
        for i in range(len(feature_importances)):
            if abs(feature_importances[i]) == max(abs(feature_importances)):
                imp = Xcol[i]

        # Create a dictionary to store the evaluation metrics
        metrics = {"Accuracy": accuracy, "Precision": precision, "Recall": recall, "F1-score": f1}
        # Plot the evaluation metrics
        plt.bar(metrics.keys(), metrics.values())
        plt.xlabel("Metrics")
        plt.ylabel("Score")
        plt.title("Model Evaluation Metrics")
        plt.savefig('pictures/{}.png'.format("Metrics"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("Metrics"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot confusion matrix
        sns.heatmap(confusionmatrix, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.savefig('pictures/{}.png'.format("confusionmatrix"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("confusionmatrix"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot feature importances
        plt.barh(Xcol, feature_importances)
        plt.xlabel("Mutual Information")
        plt.title("Feature Importances")
        plt.savefig('pictures/{}.png'.format("imp"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("imp"), 'rb').read()).decode('ascii'))
        plt.clf()

        question = classifierquestionset.render(section=1)
        text1 = classifieraccuracy.render(accuracy=round(accuracy, 3), classifiername="ridge classifier")
        text2 = text1 + classifierprecision.render(precision=precision)
        text3 = text2 + classifierrecall.render(recall=recall) + classifierf1.render(f1=f1)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)
        # dcc.Tabs([subtab1, subtab2, subtab3])

        aim = Xcol
        aim.insert(0, ycol)
        children = [html.P(question), html.Br(), dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'})]

        dash_tab_add(listTabs, 'KNeighborsClassifierStats', children)
        aim.remove(ycol)

        question = classifierquestionset.render(section=2)
        modelStory = classifierEvaluation.render(f1=round(f1, 3), accuracy=accuracy, precision=precision, recall=recall)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question), html.Br(),
                    html.P(modelStory)]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)

        question = classifierquestionset.render(section=3)
        crossvalidStory = classifiercv.render(cv=round(cv_scores.mean(), 3), cm=confusionmatrix)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question), html.Br(),
                    html.P(crossvalidStory)]
        dash_tab_add(listTabs, "Confusion Matrix and Cross-validation", children)

        question = classifierquestionset.render(section=4)
        ImpStory = classifierimp.render(imp=imp)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),
                    html.P(ImpStory)]
        dash_tab_add(listTabs, "Feature Importances", children)

        run_app(KNei_app, listTabs)

    def SVCClassifier_view(self, data, Xcol, ycol, accuracy, precision, recall, f1, confusionmatrix, cv_scores, classes,
                           total_influence, most_influential_feature, coefficients):
        _base64 = []
        svm_app, listTabs = start_app()
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Create a dictionary to store the evaluation metrics
        metrics = {"Accuracy": accuracy, "Precision": precision, "Recall": recall, "F1-score": f1}
        # Plot the evaluation metrics
        plt.bar(metrics.keys(), metrics.values())
        plt.xlabel("Metrics")
        plt.ylabel("Score")
        plt.title("Model Evaluation Metrics")
        plt.savefig('pictures/{}.png'.format("Metrics"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("Metrics"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot confusion matrix
        sns.heatmap(confusionmatrix, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.savefig('pictures/{}.png'.format("confusionmatrix"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("confusionmatrix"), 'rb').read()).decode('ascii'))
        plt.clf()

        question = classifierquestionset.render(section=1)
        text1 = classifieraccuracy.render(accuracy=round(accuracy, 3), classifiername="ridge classifier")
        text2 = text1 + classifierprecision.render(precision=precision)
        text3 = text2 + classifierrecall.render(recall=recall) + classifierf1.render(f1=f1)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        aim = Xcol
        aim.insert(0, ycol)
        children = [html.P(question), html.Br(), dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'})]
        dash_tab_add(listTabs, 'SupportVectorMachineModelStats', children)
        aim.remove(ycol)

        question = classifierquestionset.render(section=2)
        modelStory = classifierEvaluation.render(f1=round(f1, 3), accuracy=accuracy, precision=precision, recall=recall)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question), html.Br(),
                    html.P(modelStory)]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)

        question = classifierquestionset.render(section=3)
        crossvalidStory = classifiercv.render(cv=round(cv_scores.mean(), 3), cm=confusionmatrix)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question), html.Br(),
                    html.P(crossvalidStory)]
        dash_tab_add(listTabs, "Confusion Matrix and Cross-validation", children)

        question = classifierquestionset.render(section=5)
        coeffstory = classifiercoeff.render(classes=classes, Xcol=Xcol, coefficients=coefficients,
                                            total_influence=total_influence,
                                            most_influential_feature=most_influential_feature, enumerate=enumerate)
        children = [html.P(question), html.Br(),
                    html.P(coeffstory)]
        dash_tab_add(listTabs, "Feature Coefficients ", children)

        run_app(svm_app, listTabs)

    def DecisionTreeClassifier_view(self, data, Xcol, ycol, accuracy, precision, feature_importances, recall, f1,
                                    confusionmatrix, cv_scores, clf):
        _base64 = []
        dtc_app, listTabs = start_app()
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Create a dictionary to store the evaluation metrics
        metrics = {"Accuracy": accuracy, "Precision": precision, "Recall": recall, "F1-score": f1}
        # Plot the evaluation metrics
        plt.bar(metrics.keys(), metrics.values())
        plt.xlabel("Metrics")
        plt.ylabel("Score")
        plt.title("Model Evaluation Metrics")
        plt.savefig('pictures/{}.png'.format("Metrics"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("Metrics"), 'rb').read()).decode('ascii'))
        plt.clf()
        # Plot confusion matrix
        sns.heatmap(confusionmatrix, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.savefig('pictures/{}.png'.format("confusionmatrix"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("confusionmatrix"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot confusion matrix
        plt.barh(Xcol, feature_importances)
        plt.xlabel("Mutual Information")
        plt.title("Feature Importances")
        plt.savefig('pictures/{}.png'.format("imp"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("imp"), 'rb').read()).decode('ascii'))
        plt.clf()

        for i in range(len(feature_importances)):
            if abs(feature_importances[i]) == max(abs(feature_importances)):
                imp = Xcol[i]

        question = classifierquestionset.render(section=1)
        text1 = classifieraccuracy.render(accuracy=round(accuracy, 3), classifiername="decision tree classifier")
        text2 = text1 + classifierprecision.render(precision=precision)
        text3 = text2 + classifierrecall.render(recall=recall) + classifierf1.render(f1=f1)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        aim = Xcol
        aim.insert(0, ycol)
        children = [html.P(question), html.Br(), dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'})]
        dash_tab_add(listTabs, 'DecisionTreeClassifierModelStats', children)
        aim.remove(ycol)

        question = classifierquestionset.render(section=2)
        modelStory = classifierEvaluation.render(f1=round(f1, 3), accuracy=accuracy, precision=precision, recall=recall)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question), html.Br(),
                    html.P(modelStory)]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)

        question = classifierquestionset.render(section=3)
        crossvalidStory = classifiercv.render(cv=round(cv_scores.mean(), 3), cm=confusionmatrix)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question), html.Br(),
                    html.P(crossvalidStory)]
        dash_tab_add(listTabs, "Confusion Matrix and Cross-validation", children)

        question = classifierquestionset.render(section=4)
        ImpStory = classifierimp.render(imp=imp)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),
                    html.P(ImpStory)]
        dash_tab_add(listTabs, "Feature Importances", children)

        # Figure of the tree
        fig2, ax = plt.subplots(figsize=(20, 10), dpi=200)
        tree.plot_tree(clf,
                       feature_names=Xcol,
                       class_names=ycol,
                       filled=True,
                       node_ids=True);
        fig2.savefig('pictures/{}.png'.format("DT"))
        plt.clf()
        encoded_image = base64.b64encode(open("pictures/DT.png", 'rb').read()).decode('ascii')

        # # Explain of the tree
        explain = TreeExplain(clf, Xcol)
        # Text need to fix here
        tree_explain_story = explain
        question2 = DecisionTreeQuestion.render(q=2)
        children = [html.Img(src='data:image/png;base64,{}'.format(encoded_image)),
                    html.P(question2), html.Br(), html.Pre(tree_explain_story)]
        dash_tab_add(listTabs, 'Tree Explanation', children)

        run_app(dtc_app, listTabs)

    def RandomForestClassifier_view(self, data, Xcol, ycol, accuracy, precision, feature_importances, recall, f1,
                                    confusionmatrix, cv_scores):
        _base64 = []
        rfc_app, listTabs = start_app()
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        # Create a dictionary to store the evaluation metrics
        metrics = {"Accuracy": accuracy, "Precision": precision, "Recall": recall, "F1-score": f1}
        # Plot the evaluation metrics
        plt.bar(metrics.keys(), metrics.values())
        plt.xlabel("Metrics")
        plt.ylabel("Score")
        plt.title("Model Evaluation Metrics")
        plt.savefig('pictures/{}.png'.format("Metrics"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("Metrics"), 'rb').read()).decode('ascii'))
        plt.clf()
        # Plot confusion matrix
        sns.heatmap(confusionmatrix, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.savefig('pictures/{}.png'.format("confusionmatrix"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("confusionmatrix"), 'rb').read()).decode('ascii'))
        plt.clf()

        # Plot confusion matrix
        plt.barh(Xcol, feature_importances)
        plt.xlabel("Mutual Information")
        plt.title("Feature Importances")
        plt.savefig('pictures/{}.png'.format("imp"))
        _base64.append(base64.b64encode(open('pictures/{}.png'.format("imp"), 'rb').read()).decode('ascii'))
        plt.clf()

        for i in range(len(feature_importances)):
            if abs(feature_importances[i]) == max(abs(feature_importances)):
                imp = Xcol[i]

        question = classifierquestionset.render(section=1)
        text1 = classifieraccuracy.render(accuracy=round(accuracy, 3), classifiername="random forest classifier")
        text2 = text1 + classifierprecision.render(precision=precision)
        text3 = text2 + classifierrecall.render(recall=recall) + classifierf1.render(f1=f1)
        subtab1, subtab2, subtab3 = ThreeSubtab(text1, text2, text3)

        aim = Xcol
        aim.insert(0, ycol)
        children = [html.P(question), html.Br(), dcc.Tabs([subtab1, subtab2, subtab3]),
                    dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'})]
        dash_tab_add(listTabs, 'RandomForestClassifierModelStats', children)
        aim.remove(ycol)

        question = classifierquestionset.render(section=2)
        modelStory = classifierEvaluation.render(f1=round(f1, 3), accuracy=accuracy, precision=precision, recall=recall)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[0])), html.P(question), html.Br(),
                    html.P(modelStory)]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)

        question = classifierquestionset.render(section=3)
        crossvalidStory = classifiercv.render(cv=round(cv_scores.mean(), 3), cm=confusionmatrix)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[1])), html.P(question), html.Br(),
                    html.P(crossvalidStory)]
        dash_tab_add(listTabs, "Confusion Matrix and Cross-validation", children)

        question = classifierquestionset.render(section=4)
        ImpStory = classifierimp.render(imp=imp)
        children = [html.Img(src='data:image/png;base64,{}'.format(_base64[2])), html.P(question), html.Br(),
                    html.P(ImpStory)]
        dash_tab_add(listTabs, "Feature Importances", children)


        run_app(rfc_app, listTabs)

    def RidgeRegressionModel_view(self, data, Xcol, ycol, DTData, mse, mae, r2, imp, lessimp, train_r2, test_r2,
                                  train_mae, test_mae, cv_r2_scores, cv_r2_mean, mape,vif):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        for ind in DTData.index:
            ax = sns.regplot(x=ind, y=ycol, data=data)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()
        # Create Format index with file names
        _base64 = []
        for ind in DTData.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))

        rr_app, listTabs = start_app()

        fig = px.bar(DTData)

        text1 = RRA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = RRA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = RRA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="rr")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'RidgeRegressionModelStats', children)
        aim.remove(ycol)

        i = 0
        # Add to dashbord Xcol plots and data story

        for ind in DTData.index:
            question = regressionQuestion.render(xcol=ind, ycol=ycol, section=2)

            text1 = linearA2_1.render(xcol=ind, ycol=ycol, coeff=DTData['important'][ind])

            text2 = text1+' '+regressionvif.render(xcol=ind, ycol=ycol,
                                      vif_value=vif[vif['feature'] == ind]['VIF'].values[0])
            subtab1,subtab2=TwoSubtab(text1,text2)
            children = [
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])), html.P(question), html.Br(),
                dcc.Tabs([subtab1, subtab2])
            ]
            dash_tab_add(listTabs, ind, children)

            i = i + 1

        DTData['important'] = DTData['important'].round(3)
        positive_effects = DTData[DTData['important'] > 0].index.tolist()
        negative_effects = DTData[DTData['important'] < 0].index.tolist()
        no_effects = DTData[DTData['important'] == 0].index.tolist()

        question = regressionQuestion.render(ycol=ycol, section=3)
        summary = regressionSummary.render(positive_effects=positive_effects, negative_effects=negative_effects,
                                           no_effects=no_effects, imp=imp, ycol=ycol)

        children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)



        run_app(rr_app, listTabs)

    def LassoRegressionModel_view(self, data, Xcol, ycol, DTData, mse, mae, r2, imp, lessimp, train_r2, test_r2,
                                  train_mae, test_mae, cv_r2_scores, cv_r2_mean, mape,vif):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        for ind in DTData.index:
            ax = sns.regplot(x=ind, y=ycol, data=data)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()
        # Create Format index with file names
        _base64 = []
        for ind in DTData.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))

        rr_app, listTabs = start_app()

        fig = px.bar(DTData)

        text1 = LRA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = LRA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = LRA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="lr")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'LassoRegressionModelStats', children)
        aim.remove(ycol)

        i = 0
        # Add to dashbord Xcol plots and data story

        for ind in DTData.index:
            question = regressionQuestion.render(xcol=ind, ycol=ycol, section=2)

            text1 = linearA2_1.render(xcol=ind, ycol=ycol, coeff=DTData['important'][ind])

            text2 = text1+' '+regressionvif.render(xcol=ind, ycol=ycol,
                                      vif_value=vif[vif['feature'] == ind]['VIF'].values[0])
            subtab1,subtab2=TwoSubtab(text1,text2)
            children = [
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])), html.P(question), html.Br(),
                dcc.Tabs([subtab1, subtab2])
            ]
            dash_tab_add(listTabs, ind, children)

            i = i + 1

        DTData['important'] = DTData['important'].round(3)
        positive_effects = DTData[DTData['important'] > 0].index.tolist()
        negative_effects = DTData[DTData['important'] < 0].index.tolist()
        no_effects = DTData[DTData['important'] == 0].index.tolist()

        question = regressionQuestion.render(ycol=ycol, section=3)
        summary = regressionSummary.render(positive_effects=positive_effects, negative_effects=negative_effects,
                                           no_effects=no_effects, imp=imp, ycol=ycol)

        children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(rr_app, listTabs)

    def ElasticNetModel_view(self, data, Xcol, ycol, DTData, mse, mae, r2, imp, lessimp, train_r2, test_r2, train_mae,
                             test_mae, cv_r2_scores, cv_r2_mean, mape,vif):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        for ind in DTData.index:
            ax = sns.regplot(x=ind, y=ycol, data=data)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()
        # Create Format index with file names
        _base64 = []
        for ind in DTData.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))

        rr_app, listTabs = start_app()

        fig = px.bar(DTData)

        text1 = ENA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = ENA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = ENA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="en")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'ElasticNetModelStats', children)
        aim.remove(ycol)

        i = 0
        # Add to dashbord Xcol plots and data story

        for ind in DTData.index:
            question = regressionQuestion.render(xcol=ind, ycol=ycol, section=2)

            text1 = linearA2_1.render(xcol=ind, ycol=ycol, coeff=DTData['important'][ind])

            text2 = text1+' '+regressionvif.render(xcol=ind, ycol=ycol,
                                      vif_value=vif[vif['feature'] == ind]['VIF'].values[0])
            subtab1,subtab2=TwoSubtab(text1,text2)
            children = [
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])), html.P(question), html.Br(),
                dcc.Tabs([subtab1, subtab2])
            ]
            dash_tab_add(listTabs, ind, children)

            i = i + 1

        DTData['important'] = DTData['important'].round(3)
        positive_effects = DTData[DTData['important'] > 0].index.tolist()
        negative_effects = DTData[DTData['important'] < 0].index.tolist()
        no_effects = DTData[DTData['important'] == 0].index.tolist()

        question = regressionQuestion.render(ycol=ycol, section=3)
        summary = regressionSummary.render(positive_effects=positive_effects, negative_effects=negative_effects,
                                           no_effects=no_effects, imp=imp, ycol=ycol)

        children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(rr_app, listTabs)

    def LeastAngleRegressionModel_view(self, data, Xcol, ycol, DTData, mse, mae, r2, imp, lessimp, train_r2, test_r2,
                                       train_mae, test_mae, cv_r2_scores, cv_r2_mean, mape,vif):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        for ind in DTData.index:
            ax = sns.regplot(x=ind, y=ycol, data=data)
            plt.savefig('pictures/{}.png'.format(ind))
            plt.clf()
        # Create Format index with file names
        _base64 = []
        for ind in DTData.index:
            _base64.append(base64.b64encode(open('pictures/{}.png'.format(ind), 'rb').read()).decode('ascii'))

        rr_app, listTabs = start_app()

        fig = px.bar(DTData)

        text1 = LARA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = LARA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = LARA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="lar")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'LeastAngleRegressionModelStats', children)
        aim.remove(ycol)

        i = 0
        # Add to dashbord Xcol plots and data story

        for ind in DTData.index:
            question = regressionQuestion.render(xcol=ind, ycol=ycol, section=2)

            text1 = linearA2_1.render(xcol=ind, ycol=ycol, coeff=DTData['important'][ind])

            text2 = text1+' '+regressionvif.render(xcol=ind, ycol=ycol,
                                      vif_value=vif[vif['feature'] == ind]['VIF'].values[0])
            subtab1,subtab2=TwoSubtab(text1,text2)
            children = [
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])), html.P(question), html.Br(),
                dcc.Tabs([subtab1, subtab2])
            ]
            dash_tab_add(listTabs, ind, children)

            i = i + 1

        DTData['important'] = DTData['important'].round(3)
        positive_effects = DTData[DTData['important'] > 0].index.tolist()
        negative_effects = DTData[DTData['important'] < 0].index.tolist()
        no_effects = DTData[DTData['important'] == 0].index.tolist()

        question = regressionQuestion.render(ycol=ycol, section=3)
        summary = regressionSummary.render(positive_effects=positive_effects, negative_effects=negative_effects,
                                           no_effects=no_effects, imp=imp, ycol=ycol)

        children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(rr_app, listTabs)

    def AdaBoostRegressionModel_view(self, data, Xcol, ycol, DTData, mse, mae, r2, imp, lessimp, train_r2, test_r2,
                                     train_mae, test_mae, cv_r2_scores, cv_r2_mean, mape):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')

        rr_app, listTabs = start_app()

        fig = px.bar(DTData)

        text1 = ADAA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = ADAA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = ADAA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="ada")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'AdaBoostRegressionModelStats', children)
        aim.remove(ycol)

        DTData['important'] = DTData['important'].round(3)
        positive_effects = DTData[DTData['important'] > 0].index.tolist()
        negative_effects = DTData[DTData['important'] < 0].index.tolist()
        no_effects = DTData[DTData['important'] == 0].index.tolist()

        question = regressionQuestion.render(ycol=ycol, section=3)
        summary = regressionSummary.render(positive_effects=positive_effects, negative_effects=negative_effects,
                                           no_effects=no_effects, imp=imp, ycol=ycol)

        children = [dcc.Graph(figure=fig), html.P(question), html.Br(), html.P(summary)]
        dash_tab_add(listTabs, 'Summary', children)

        run_app(rr_app, listTabs)

    def KNeighborsRegressionModel_view(self, data, Xcol, ycol, model, mse, mae, r2, train_r2, test_r2,
                                       train_mae, test_mae, cv_r2_scores, cv_r2_mean, mape):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')

        rr_app, listTabs = start_app()

        text1 = KNRA1_1.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol)
        rmse = mse ** (1 / 2.0)
        text2 = KNRA1_2.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape)
        text3 = KNRA1_3.render(r2=r2, indeNum=np.size(Xcol), Xcol=Xcol, ycol=ycol, mape=mape, mse=mse, rmse=rmse,
                               mae=mae)

        question1 = regressionQuestion.render(section=1, m="knr")
        question2 = question1+' '+linearQuestion.render(section=4)
        question3 = question2+' '+linearQuestion.render(section=5)+' '+linearQuestion.render(section=6)
        subtab1, subtab2, subtab3 = ThreeSubQA(question1,question2,question3,text1, text2, text3)

        # intro = MicroLexicalization(intro)
        aim = Xcol
        aim.insert(0, ycol)
        children = [dash_table.DataTable(data[aim].to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          data[aim].columns],
                                         style_table={'height': '400px',
                                                      'overflowY': 'auto'}), dcc.Tabs([subtab1, subtab2, subtab3]),
                    ]
        dash_tab_add(listTabs, 'KNeighborsRegressionModelStats', children)
        aim.remove(ycol)

        run_app(rr_app, listTabs)

    def FindBestRegression_view(self,data,Xcol,ycol,comapre_results,fitstory,modelcomparestory,p_values_df):
        _base64 = []
        BR_app, listTabs = start_app()

        question='Which regression model might be the best fit for the dataset?'

        children = [html.P(html.B(question)),html.P(modelcomparestory),html.Br(),html.P(fitstory),
                    dash_table.DataTable(comapre_results.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          comapre_results.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        create_scatter_plots(data, Xcol, ycol)
        for ind in Xcol:
            _base64.append(base64.b64encode(open(f'pictures/{ind}.png', 'rb').read()).decode('ascii'))

        # Add tabs with scatter plots
        for i, ind in enumerate(Xcol):
            question='Does the independent variable have a significant effect on the dependent variable?'
            p_value = p_values_df[p_values_df['Feature'] == ind]['P-Value'].values[0]
            p_value_story=pvaluesummary.render(p_value=p_value)
            tab_content = [html.P(html.B(question)),html.P(p_value_story),
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])),
                html.P(f'Scatter plot of {ind} vs {ycol}')
            ]
            dash_tab_add(listTabs, f'{ind} vs {ycol}', tab_content)

        run_app(BR_app, listTabs)

    def FindBestClassifier_view(self,data,Xcol,ycol,comapre_results,fitstory,modelcomparestory,importance,pycaretname,modelname):
        if not os.path.exists('pictures'):
            os.makedirs('pictures')
        _base64 = []
        BR_app, listTabs = start_app()

        question='Which classifier model might be the best fit for the dataset?'

        children = [html.P(html.B(question)),html.P(modelcomparestory),html.Br(),html.P(fitstory),
                    dash_table.DataTable(comapre_results.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          comapre_results.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "Model Evaluation Metrics", children)

        create_scatter_plots(data, Xcol, ycol)
        for ind in Xcol:
            _base64.append(base64.b64encode(open(f'pictures/{ind}.png', 'rb').read()).decode('ascii'))

        # Add tabs with scatter plots
        for i, ind in enumerate(Xcol):
            # question='Does the independent variable have a significant effect on the dependent variable?'
            # p_value = p_values_df[p_values_df['Feature'] == ind]['P-Value'].values[0]
            # p_value_story=pvaluesummary.render(p_value=p_value)
            max_feature = importance.iloc[0]['Feature']
            max_value = importance.iloc[0]['Value']
            # print(max_feature)
            # print(max_value)
            if pycaretname in ['lr', 'lda', 'ridge', 'svm']:
                question = 'On average, what is the impact of this independent variable on the importance of the classification result?'
                story=pycaretclassificoefimp1.render(modelname=modelname,coef=importance[importance['Feature'] == ind]['Value'].values[0],max_value=max_value)
                # story = importance[importance['Feature'] == ind]['Value'].values[0]
            elif pycaretname in ['rf', 'et', 'gbc', 'xgboost', 'lightgbm', 'catboost', 'ada', 'dt']:
                question = 'What is the impact of this independent variable on the importance of the classification result?'
                #story = importance[importance['Feature'] == ind]['Value'].values[0]
                story=pycaretclassificoefimp2.render(modelname=modelname,imp=importance[importance['Feature'] == ind]['Value'].values[0],max_value=max_value)

            tab_content = [html.P(html.B(question)),html.P(story),
                html.Img(src='data:image/png;base64,{}'.format(_base64[i])),
                html.P(f'Scatter plot of {ind} vs {ycol}')
            ]
            dash_tab_add(listTabs, f'{ind} vs {ycol}', tab_content)

        run_app(BR_app, listTabs)


class DocumentplanningNoDashboard:

    def dataset_description(self,df):
        corr_matrix = df.corr()
        description=''
        for i in range(len(corr_matrix.columns)):
            for j in range(i + 1, len(corr_matrix.columns)):
                var1 = corr_matrix.columns[i]
                var2 = corr_matrix.columns[j]
                correlation = corr_matrix.iloc[i, j]
                description = description+correlation_story.render(var1=var1, var2=var2, correlation=correlation)+"\n"
        print(description)
        sns.heatmap(df.corr(), annot=True, cmap='coolwarm')
        plt.show()
        return (description)

    def simple_dataset_description(self, df):
        corr_matrix = df.corr()
        positive_pairs = []
        negative_pairs = []
        for i in range(len(corr_matrix.columns)):
            for j in range(i + 1, len(corr_matrix.columns)):
                var1 = corr_matrix.columns[i]
                var2 = corr_matrix.columns[j]
                correlation = corr_matrix.iloc[i, j]
                if correlation > 0:
                    positive_pairs.append((var1, var2))
                elif correlation < 0:
                    negative_pairs.append((var1, var2))
        description = correlation_story_2.render(positive_pairs=positive_pairs, negative_pairs=negative_pairs)
        print(description)
        sns.heatmap(df.corr(), annot=True, cmap='coolwarm')
        plt.show()
        return (description)

    def dataset_information(self,df):


        missing_values = df.isnull().sum()
        unique_values = df.nunique()
        columns = df.columns.tolist()
        description = datasetinformation.render(columns=columns, missing_values=missing_values.to_dict(),
                                      unique_values=unique_values.to_dict())
        print(description)

        df.hist(bins=30, figsize=(20, 15))
        plt.show()
        return (description)

    def CP_general_description(self,ycol,last_X,last_X2,last_y,difference,percentage_change,max_value,max_y_X):

        summary = piecewiseCP2.render(Xend=last_X, ycol=ycol,yend=last_y, diff=difference,percentagechange=percentage_change,Xlast=last_X2,ymax=max_value,Xmax=max_y_X)
        # print(summary)
        return (summary)
    def CP_piecewise_linear(self,data, Xcol, ycol, slopes,breaks):
        def merge_segments(breaks, slopes):
            merged_breaks = [breaks[0]]
            merged_slopes = []
            current_slope = slopes[0]
            current_start = breaks[0]
            for i in range(1, len(slopes)):
                if (slopes[i] > 0 and current_slope > 0) or (slopes[i] < 0 and current_slope < 0) or (
                        slopes[i] == 0 and current_slope == 0):
                    continue
                else:
                    merged_breaks.append(breaks[i])
                    merged_slopes.append(current_slope)
                    current_slope = slopes[i]
                    current_start = breaks[i]
            merged_breaks.append(breaks[-1])
            merged_slopes.append(current_slope)
            output_slopes = [1 if s > 0 else -1 if s < 0 else 0 for s in merged_slopes]
            return merged_breaks, output_slopes

        merged_breaks, output_slopes = merge_segments(breaks, slopes)
        rounded_breaks = [math.ceil(value) for value in merged_breaks]
        # print(rounded_breaks,output_slopes)
        summary = piecewiseCP1.render(xcol=Xcol, ycol=ycol, merged_breaks=rounded_breaks,
                                           output_slopes=output_slopes)
        # print(summary)
        return (summary)

    def DRD_piecewise_linear(self,data, Xcol, ycol,model, slopes, segment_points, segment_values, max_slope_info,breaks):
        def merge_segments(breaks, slopes):
            merged_breaks = [breaks[0]]
            merged_slopes = []
            current_slope = slopes[0]
            current_start = breaks[0]
            for i in range(1, len(slopes)):
                if (slopes[i] > 0 and current_slope > 0) or (slopes[i] < 0 and current_slope < 0) or (
                        slopes[i] == 0 and current_slope == 0):
                    continue
                else:
                    merged_breaks.append(breaks[i])
                    merged_slopes.append(current_slope)
                    current_slope = slopes[i]
                    current_start = breaks[i]
            merged_breaks.append(breaks[-1])
            merged_slopes.append(current_slope)
            output_slopes = [1 if s > 0 else -1 if s < 0 else 0 for s in merged_slopes]
            return merged_breaks, output_slopes

        merged_breaks, output_slopes = merge_segments(breaks, slopes)
        rounded_breaks = [math.ceil(value) for value in merged_breaks]

        # print(rounded_breaks,output_slopes)
        summary = piecewiseDRD1.render(xcol=Xcol, ycol=ycol,merged_breaks=rounded_breaks, output_slopes=output_slopes,startPoint=int(max_slope_info['start_end_points'][0]),endPoint=max_slope_info['start_end_points'][1],slope=max_slope_info['slope'])
        # print(summary)
        return (summary)

class AutoFindBestModel():
    def __init__(self):
        pass

    def model_compare(self, modelname, modeldetail, selected_criterion,n=1):
        modelcomparestory = automodelcompare1.render(best=modelname, detail=modeldetail, n_select=n, sort=selected_criterion)
        return modelcomparestory

    def pycaret_model_summary_view(self,imp_var, r2, mape,target,modelname):
        story1 = pycaretmodelfit.render(r2=r2, mape=mape,modelname=modelname)
        # story2 = pycaretimp.render(imp=imp_var, target=target, imp_pos_ave=imp_pos_ave,
        #                            imp_pos_value_ave=imp_pos_value_ave,
        #                            imp_neg_ave=imp_neg_ave, imp_neg_value_ave=imp_neg_value_ave)

        return (story1)

    def pycaret_classificationmodel_summary_view(self,imp, Accuracy, AUC,target,modelname):
        story1 = pycaretclassificationfit.render(accuracy=Accuracy, auc=AUC, modelname=modelname)
        return (story1)

class special_view_for_ACCCP:
    def register_question1_view(self,register_dataset, per1000inCity_col, diff, table_col, label, app, listTabs):
        registerstory = "The data from local comparators features in the Child Protection Register (CPR) report prepared quarterly. "
        i = 0
        for ind in per1000inCity_col:
            reslut = register_story.render(Xcol=ind, minX=min(register_dataset[ind]), maxX=max(register_dataset[ind]),
                                           diff=diff[i])
            registerstory = registerstory + reslut
            i = i + 1
        children = [html.P(registerstory),
                    dash_table.DataTable(register_dataset.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          register_dataset.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "question1-2", children)
        return (app,listTabs,registerstory)

    def riskfactor_question1_view(self,dataset, max_factor, same_factor, label, cityname, app, listTabs):
        riskstory = risk_factor_story.render(indeNum=(np.size(max_factor)), max_factor=max_factor,
                                             same_factor=same_factor,
                                             cityname=cityname)
        children = [html.P(riskstory),
                    dash_table.DataTable(dataset.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          dataset.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "question1-3", children)
        return (app, listTabs,riskstory)

    def re_register_question4_view(self,register_dataset, national_average_reregistration, reregister_lastyear, period,
                                   table_col, label, app, listTabs):
        reregisterstory = reregister_story.render(nar=national_average_reregistration, rrly=reregister_lastyear,
                                                  time=period)
        children = [html.P(reregisterstory),
                    dash_table.DataTable(register_dataset.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          register_dataset.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "question4", children)
        return (app, listTabs,reregisterstory)

    def remain_time_question5_view(self,remain_data, zero_lastdata, label, app, listTabs):
        remainstory = remain_story.render(
            zl=zero_lastdata)  # It can do more if I know the rule of answering this question
        children = [html.P(remainstory),
                    dash_table.DataTable(remain_data.to_dict('records'),
                                         [{"name": i, "id": i} for i in
                                          remain_data.columns],style_table={'height': '400px', 'overflowY': 'auto'})    ]
        dash_tab_add(listTabs, "question5", children)
        return (app, listTabs,remainstory)

    def enquiries_question6_view(self,ACmean, ASmean, MTmean, ACdata, ASdata, MTdata, period, label, app, listTabs):
        enquiriesstory = enquiries_story.render(indeNum=(np.size(period)), ACM=ACmean, ASM=ASmean, MTM=MTmean,
                                                ACE=ACdata,
                                                ASE=ASdata,
                                                MTE=MTdata, period=period)
        children = [html.P(enquiriesstory) ]
        dash_tab_add(listTabs, "question6", children)
        return (app, listTabs,enquiriesstory)

    def national_data_view(self, change_percentage, initial_year_value, final_year_value, shirechange_percentage,
                           Moraychange_percentage, IPchange_percentage):
        nationalstory = national_data_compare_story.render(change_percentage=change_percentage,
                                                           initial_year_value=initial_year_value,
                                                           final_year_value=final_year_value,
                                                           shirechange_percentage=shirechange_percentage,
                                                           Moraychange_percentage=Moraychange_percentage,
                                                           IPchange_percentage=IPchange_percentage)
        children = [html.P(nationalstory)]
        return (nationalstory)
        # dash_tab_add(listTabs, "question1-1", children)
        # return (app, listTabs,enquiriesstory)

    def timescales_description(self, df, colname1, colname2, colname3, colname4, colname5, colname6, colname7,
                               colname8, colname9, colname10, colname11, colname12, colname13,
                               colname14, colname15, colname16, colname17, colname18, colname19,
                               colname20, colname21, colname22, colname23, colname24):
        casenumber1 = df[colname1][0]
        casenumber2 = df[colname2][0]
        casenumber3 = df[colname3][0]
        percentage3 = round(casenumber3 / (casenumber3 + df[colname24][0]) * 100)
        percentage3_2 = df[colname21][0]
        percentage4_jun = df[colname4][0]
        percentage4_sep = df[colname5][0]
        percentage4_dec = df[colname6][0]

        casenumber5 = df[colname7][0]
        percentage5 = round(casenumber5 / (casenumber5 + df[colname8][0]) * 100)

        casenumber6_1 = df[colname22][0]
        casenumber6_2 = df[colname23][0]
        percentage6 = round(casenumber6_1 / (casenumber6_1 + casenumber6_2) * 100)

        casenumber7 = df[colname9][0]
        percentage7 = round(casenumber7 / (casenumber7 + df[colname10][0]) * 100)

        casenumber8 = df[colname11][0]
        percentage8 = round(casenumber8 / (casenumber8 + df[colname12][0]) * 100)

        casenumber9 = df[colname13][0]
        percentage9 = round(casenumber9 / (casenumber9 + df[colname14][0]) * 100)

        casenumber10 = df[colname15][0]
        percentage10 = round(casenumber10 / (casenumber10 + df[colname16][0]) * 100)

        casenumber11 = df[colname17][0]
        percentage11 = round(casenumber11 / (casenumber11 + df[colname18][0]) * 100)

        casenumber12 = df[colname18][0]

        casenumber13 = df[colname19][0]
        percentage13 = round(casenumber13 / (casenumber13 + df[colname20][0]) * 100)

        casenumber14 = df[colname20][0]
        timescalesstory = timescales_story.render(colname1=colname1, casenumber1=casenumber1,
                                                  colname2=colname2, casenumber2=casenumber2,
                                                  colname3=colname3, casenumber3=casenumber3, percentage3=percentage3,
                                                  colname4=colname4, percentage4_jun=percentage4_jun,
                                                  colname5=colname5, percentage4_sep=percentage4_sep,
                                                  colname6=colname6, percentage4_dec=percentage4_dec,
                                                  colname7=colname7, casenumber5=casenumber5, percentage5=percentage5,
                                                  casenumber6=casenumber6_1, percentage6=percentage6,
                                                  colname8=colname9, casenumber7=casenumber7, percentage7=percentage7,
                                                  colname9=colname11, casenumber8=casenumber8, percentage8=percentage8,
                                                  colname10=colname13, casenumber9=casenumber9, percentage9=percentage9,
                                                  colname11=colname15, casenumber10=casenumber10,
                                                  percentage10=percentage10,
                                                  colname12=colname17, casenumber11=casenumber11,
                                                  percentage11=percentage11,
                                                  casenumber12=casenumber12,
                                                  colname13=colname19, casenumber13=casenumber13,
                                                  percentage13=percentage13,
                                                  casenumber14=casenumber14, percentage3_2=percentage3_2)
        return timescalesstory

    def SCRA_description(self,total_1,total_2,point1):
        SCRAstory=SCRA_story.render(total_1=total_1,total_2=total_2,point1=point1)
        return SCRAstory

    def unfinished_report(self,app_name, listTabs,story1,story2,story3,story4,story5,template_path,output_path='./output.docx'):
        #unfinishedreport=unfinished_report_template.render(story1=story1,story2=story2,story3=story3,story4=story4,story5=story5)
        # template_path = './templates/report_template.docx'
        doc = Document(template_path)

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        template_str = '\n'.join(full_text)

        jinja_template = Template(template_str)
        rendered_str = jinja_template.render(story1=story1,story2=story2,story3=story3,story4=story4,story5=story5)
        # print(story1)
        # replacements = {
        #     "{{story1}}": story1,
        #     "{{story2}}": story2,
        #     "{{story3}}": story3,
        #     "{{story4}}": story4,
        #     "{{story5}}": story5
        # }
        #
        # fill_template(template_path, output_path, replacements)
        rendered_str=rendered_str+"New document saved to {output_path}"
        # print(f"New document saved to {output_path}")
        # new_doc = Document()
        # for line in rendered_str.split('\n'):
        #     new_doc.add_paragraph(line)
        # output_path = './Rendered_Data_Report.docx'
        # new_doc.save(output_path)
        #
        # download_link = html.A("Click here to download a report template that answers the previous questions.", href=f'/download/{os.path.basename(output_path)}', target="_blank")
        #
        # children = [download_link]

        children = [html.Pre(rendered_str) ]
        dash_tab_add(listTabs, "report template", children)
        return (app_name, listTabs)



def dependentcompare_view(Xcolname, begin, end, ycolname1, ycolname2, magnification1, magnification2, X, X1, X2):

    return(dc1.render(Xcol=Xcolname, begin=begin, end=end, loopnum=end, y1name=ycolname1, y2name=ycolname2,
                     magnification1=magnification1,
                     magnification2=magnification2, X=X, X1=X1, X2=X2))


def batchprovessing_view1(m, Xcolname, X1, X2, y, allincrease, alldecrease, category_name, ycolnames, begin, end):
    story = (bp1.render(mode=m, Xcol=Xcolname, X1=X1, allincrease=allincrease, alldecrease=alldecrease,
                        category_name=category_name)) + "\n"
    for i in range(np.size(ycolnames) - 1):
        ycolname = ycolnames[i]
        ydata = y[ycolname]
        y1 = ydata[begin]
        y2 = ydata[end]
        story = story + '\n'+bp2.render(mode=m, ycol=ycolname, y1=y1, y2=y2, X1=X1, X2=X2, mag=0)
    return (story)


def batchprovessing_view2(m, Xcolname, X1, allincrease, alldecrease, category_name, total, ycolnames, y, point):
    story = (bp1.render(mode=m, Xcol=Xcolname, X1=X1, allincrease=False, alldecrease=False,
                        category_name=category_name)) + "\n"
    for i in range(np.size(ycolnames) - 1):
        ycolname = ycolnames[i]
        ydata = y[ycolname]
        y1 = ydata[point]
        mag = np.round(y1 / total, 2)
        story = story + bp2.render(mode=m, ycol=ycolname, y1=y1, y2=0, X1=0, X2=0, mag=mag)
    print(story)

def independenttwopointcompare_view(Xcolname, point, ycolname1, ycolname2, X, y1, y2, mode, mag):
    return (idtpc.render(Xcol=Xcolname, point=point, y1name=ycolname1, y2name=ycolname2, X=X, y1=y1, y2=y2,
                       mode=mode, mag=mag))


def toptwo_view(top_two_age_groups):
    template = Template(
        "In 2020, {{ age1_percentage }}% of all drug-related deaths were of people aged {{ age1 }}, "
        "and followed by {{ age2_percentage }}% between {{ age2 }}."
    )

    result = template.render(
        age1_percentage=round(top_two_age_groups[0][1], 2),
        age1=top_two_age_groups[0][0].replace('Deaths ', ''),
        age2_percentage=round(top_two_age_groups[1][1], 2),
        age2=top_two_age_groups[1][0].replace('Deaths ', '')
    )
    return (result)